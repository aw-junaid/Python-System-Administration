#!/usr/bin/env python3
"""
word_automation.py
===================
Automate Microsoft Word document generation with python-docx, including
a simple mail-merge: one recipient's data merged into a letter template
per row of a CSV, saved as separate .docx files.

Usage
-----
    # Mail merge (default mode) using a CSV of recipients
    python word_automation.py --mode mail-merge --recipients recipients.csv --output-dir merged_letters

    # Generate a single standalone report/document
    python word_automation.py --mode report --output report.docx

If --recipients is omitted in mail-merge mode, a sample CSV is created
first so the script runs end-to-end.

Expected output
----------------
- mail-merge mode: one .docx letter per row in the CSV, written to
  --output-dir, with placeholders like {{name}} replaced by that row's
  values.
- report mode: a single formatted .docx with a title, headings,
  a paragraph, and a table, written to --output.

Requirements
------------
    pip install python-docx
"""

import argparse
import csv
import os

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


LETTER_TEMPLATE = (
    "Dear {{name}},\n\n"
    "Thank you for your recent purchase of {{product}}. Your order "
    "{{order_id}} has been confirmed and will ship to {{city}} shortly.\n\n"
    "Warm regards,\nCustomer Success Team"
)


def create_sample_recipients(path: str) -> None:
    content = (
        "name,product,order_id,city\n"
        "Hassan Raza,Wireless Mouse,ORD-2001,Lahore\n"
        "Ayesha Noor,Mechanical Keyboard,ORD-2002,Karachi\n"
        "Bilal Farooq,USB-C Hub,ORD-2003,Islamabad\n"
    )
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"[info] No --recipients given, created a sample CSV at: {path}")


def fill_template(template: str, row: dict) -> str:
    text = template
    for key, value in row.items():
        text = text.replace("{{" + key + "}}", str(value))
    return text


def mail_merge(recipients_csv: str, output_dir: str) -> None:
    if not os.path.exists(recipients_csv):
        print(f"[error] File not found: {recipients_csv}")
        return

    os.makedirs(output_dir, exist_ok=True)
    files_written = []

    with open(recipients_csv, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            merged_text = fill_template(LETTER_TEMPLATE, row)

            doc = Document()
            for paragraph_text in merged_text.split("\n\n"):
                p = doc.add_paragraph()
                run = p.add_run(paragraph_text)
                run.font.size = Pt(11)

            safe_name = "".join(c if c.isalnum() else "_" for c in row.get("name", "recipient"))
            out_path = os.path.join(output_dir, f"letter_{safe_name}.docx")
            doc.save(out_path)
            files_written.append(out_path)

    print(f"[success] Mail merge complete. {len(files_written)} letter(s) written to '{output_dir}':")
    for fp in files_written:
        print(f"          {fp}")


def generate_report(output: str) -> None:
    doc = Document()

    title = doc.add_heading("Monthly Operations Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Overview", level=1)
    doc.add_paragraph(
        "This report summarizes key operational metrics for the period, "
        "generated automatically via python-docx."
    )

    doc.add_heading("Key Metrics", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Metric"
    hdr_cells[1].text = "This Month"
    hdr_cells[2].text = "Last Month"

    sample_metrics = [
        ("Orders Fulfilled", "1,204", "1,110"),
        ("Avg. Delivery Time (days)", "2.3", "2.6"),
        ("Customer Satisfaction (%)", "94", "91"),
    ]
    for metric, this_month, last_month in sample_metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = this_month
        row_cells[2].text = last_month

    doc.add_paragraph()
    doc.add_paragraph("Generated automatically. No manual edits applied.").italic = True

    doc.save(output)
    print(f"[success] Report document written to: {output}")


def main():
    parser = argparse.ArgumentParser(description="Automate Word document generation and mail merge with python-docx.")
    parser.add_argument("--mode", choices=["mail-merge", "report"], default="mail-merge", help="Which automation to run")
    parser.add_argument("--recipients", default=None, help="CSV file of recipients for mail-merge mode")
    parser.add_argument("--output-dir", default="merged_letters", help="Output directory for mail-merge mode")
    parser.add_argument("--output", default="report.docx", help="Output .docx path for report mode")
    args = parser.parse_args()

    if args.mode == "mail-merge":
        recipients = args.recipients
        if recipients is None:
            recipients = "recipients.csv"
            create_sample_recipients(recipients)
        mail_merge(recipients, args.output_dir)
    else:
        generate_report(args.output)


if __name__ == "__main__":
    main()
